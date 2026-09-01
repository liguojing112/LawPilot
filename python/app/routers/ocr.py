"""OCR 文本提取服务 — PaddleOCR + pdfplumber + python-docx"""
import os

from fastapi import APIRouter
from pydantic import BaseModel

router = APIRouter(prefix="/ocr", tags=["ocr"])


class ExtractRequest(BaseModel):
    file_path: str
    file_type: str  # pdf / png / jpg / jpeg / doc / docx


class ExtractResponse(BaseModel):
    text: str
    page_count: int = 1


def _flatten_ocr_texts(result) -> list[str]:
    """兼容 PaddleOCR 2.x（[poly, (text, score)]）与 3.x（dict 含 rec_texts）返回结构"""
    texts: list[str] = []

    def _extract(item):
        if isinstance(item, dict):
            for t in item.get("rec_texts", []):
                texts.append(str(t))
        elif isinstance(item, (list, tuple)):
            for line in item:
                if isinstance(line, (list, tuple)) and len(line) >= 2:
                    t = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                    texts.append(t)

    if isinstance(result, dict):
        _extract(result)
    elif isinstance(result, (list, tuple)):
        for item in result:
            _extract(item)
    return texts


def _extract_image(file_path: str) -> tuple[str, int]:
    """使用 PaddleOCR 提取图片文本"""
    try:
        from paddleocr import PaddleOCR

        try:
            ocr = PaddleOCR(lang="ch")
        except TypeError:
            ocr = PaddleOCR(lang="ch", show_log=False)

        try:
            result = ocr.ocr(file_path)
        except (AttributeError, TypeError):
            result = ocr.predict(file_path)

        return "\n".join(_flatten_ocr_texts(result)), 1
    except ImportError:
        return "[PaddleOCR 未安装，请运行 pip install paddleocr paddlepaddle]", 1
    except Exception as e:
        raise RuntimeError(f"OCR 处理失败: {e}") from e


def _ocr_pdf_pages(file_path: str, page_indices: list[int]) -> dict[int, str]:
    """把 PDF 指定页渲染成图片逐页 OCR（扫描件降级路径），返回 页下标->文本"""
    import shutil
    import tempfile

    import pypdfium2 as pdfium

    results: dict[int, str] = {}
    doc = pdfium.PdfDocument(file_path)
    tmpdir = tempfile.mkdtemp(prefix="lawpilot_ocr_")
    try:
        for i in page_indices:
            img_path = os.path.join(tmpdir, f"page_{i}.png")
            try:
                page = doc[i]
                pil_image = page.render(scale=2.0).to_pil()
                pil_image.save(img_path)
                text, _ = _extract_image(img_path)
                results[i] = text
            except Exception:
                results[i] = ""
    finally:
        doc.close()
        shutil.rmtree(tmpdir, ignore_errors=True)
    return results


def _extract_pdf(file_path: str) -> tuple[str, int]:
    """PDF 提取：优先各页文本层；无文本层的扫描页渲染成图片 OCR（支持全部页）"""
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            n = len(pdf.pages)
            page_texts = [(page.extract_text() or "").strip() for page in pdf.pages]

        ocr_needed = [i for i, t in enumerate(page_texts) if not t]
        if ocr_needed:
            try:
                ocr_map = _ocr_pdf_pages(file_path, ocr_needed)
            except Exception:
                # pypdfium2 不可用/渲染失败：退回旧方案（整份 PDF 交给 PaddleOCR，通常只处理第一页）
                first, _ = _extract_image(file_path)
                ocr_map = {i: (first if i == ocr_needed[0] else "") for i in ocr_needed}
            for i in ocr_needed:
                page_texts[i] = ocr_map.get(i, "")

        text = "\n\n".join(t for t in page_texts if t)
        return text, n
    except ImportError:
        return "[pdfplumber 未安装]", 1
    except Exception as e:
        raise RuntimeError(f"PDF 提取失败: {e}") from e


def _extract_docx(file_path: str) -> tuple[str, int]:
    """使用 python-docx 提取 Word 文档文本"""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs), 1
    except ImportError:
        return "[python-docx 未安装]", 1
    except Exception as e:
        raise RuntimeError(f"Word 文档提取失败: {e}") from e


@router.post("/extract", response_model=ExtractResponse)
async def extract_text(req: ExtractRequest):
    """从文件中提取文本"""
    if not os.path.exists(req.file_path):
        return ExtractResponse(text="", page_count=0)

    file_type = req.file_type.lower()

    if file_type in ("png", "jpg", "jpeg", "tiff", "tif"):
        text, page_count = _extract_image(req.file_path)
    elif file_type == "pdf":
        text, page_count = _extract_pdf(req.file_path)
    elif file_type in ("doc", "docx"):
        text, page_count = _extract_docx(req.file_path)
    else:
        # 尝试作为文本文件直接读取
        with open(req.file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        page_count = 1

    return ExtractResponse(text=text, page_count=page_count)
